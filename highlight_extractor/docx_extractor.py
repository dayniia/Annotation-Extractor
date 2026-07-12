"""
docx_extractor.py — DOCX highlight extractor using python-docx.

Milestone 3:
- Iterate paragraphs and runs.
- Capture highlight_color (WD_COLOR_INDEX) name string.
- Merge adjacent runs with the same color into one highlight.
- Record paragraph index (and heading context) as location.
- Optionally capture comments anchored to highlighted text.
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Union

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
except ImportError as exc:
    raise ImportError(
        "python-docx is required for DOCX extraction. Install it with:\n"
        "  pip install python-docx"
    ) from exc

from .schema import Highlight


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _color_name(color_index: WD_COLOR_INDEX) -> str:
    """Return the lowercase canonical palette name for a WD_COLOR_INDEX value."""
    try:
        return color_index.name.lower()
    except AttributeError:
        return str(color_index).lower()


def _heading_context(doc: "Document", para_idx: int) -> Union[str, None]:
    """
    Walk backwards from para_idx to find the nearest heading paragraph.
    Returns the heading text, or None if none found before para 0.
    """
    paragraphs = doc.paragraphs
    for i in range(para_idx - 1, -1, -1):
        style_name = paragraphs[i].style.name or ""
        if style_name.startswith("Heading"):
            return paragraphs[i].text.strip() or None
    return None


# ---------------------------------------------------------------------------
# Public extractor
# ---------------------------------------------------------------------------

def extract(file_path: Union[str, Path], include_formatting: bool = False) -> list[Highlight]:
    """
    Extract all highlighted runs, comments, and (optionally) formatting from a DOCX document.

    Parameters
    ----------
    file_path:
        Path to the .docx file.
    include_formatting:
        If True, also extracts underlines and strikeouts as annotations.

    Returns
    -------
    list[Highlight]
        Highlights in reading order (paragraph index order).

    Raises
    ------
    FileNotFoundError
        If the file does not exist.
    ValueError
        If the file is not a .docx or cannot be opened.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if file_path.suffix.lower() not in (".docx",):
        raise ValueError(f"Expected a .docx file, got: {file_path.suffix!r}")

    try:
        doc = Document(str(file_path))
    except Exception as exc:
        raise ValueError(f"Could not open DOCX: {exc}") from exc

    highlights: list[Highlight] = []
    order = 0

    # 1. Parse comments part to get metadata (text, date)
    comments_meta = {}
    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break
            
    if comments_part:
        for c in comments_part.element.findall(".//w:comment", namespaces=comments_part.element.nsmap):
            c_id = c.get(qn('w:id'))
            date = c.get(qn('w:date'), None)
            text = "".join(t.text for t in c.findall(".//w:t", namespaces=c.nsmap) if t.text)
            comments_meta[c_id] = {"date": date, "note": text, "text_parts": [], "location": ""}

    # 2. Iterate paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        heading = _heading_context(doc, para_idx)
        location = f"paragraph {para_idx + 1}"
        if heading:
            location += f" (under: {heading})"

        segments: list[dict] = []  # {"color": str, "text": str, "type": str}
        active_comments = set()

        # We must iterate over XML elements directly to correctly sequence runs and comment ranges
        for child in para._element:
            tag = child.tag
            
            if tag == qn('w:commentRangeStart'):
                c_id = child.get(qn('w:id'))
                if c_id in comments_meta:
                    active_comments.add(c_id)
                    if not comments_meta[c_id]["location"]:
                        comments_meta[c_id]["location"] = location
                        
            elif tag == qn('w:commentRangeEnd'):
                c_id = child.get(qn('w:id'))
                if c_id in active_comments:
                    active_comments.remove(c_id)
                    
            elif tag == qn('w:r'):
                # Wrap it in a python-docx Run object to easily read properties
                from docx.text.run import Run
                run = Run(child, para)
                run_text = run.text
                if not run_text:
                    continue

                # Add text to active comments
                if active_comments:
                    for c_id in active_comments:
                        comments_meta[c_id]["text_parts"].append(run_text)

                # Check highlight
                hc = run.font.highlight_color
                if hc is not None:
                    color_str = _color_name(hc)
                    if segments and segments[-1].get("_open") and segments[-1]["color"] == color_str and segments[-1]["type"] == "highlight":
                        segments[-1]["text"] += run_text
                    else:
                        if segments and segments[-1].get("_open"):
                            segments[-1]["_open"] = False
                        segments.append({"color": color_str, "text": run_text, "type": "highlight", "_open": True})
                else:
                    if segments and segments[-1].get("_open"):
                        segments[-1]["_open"] = False
                
                # Check formatting if enabled
                if include_formatting:
                    if run.underline:
                        segments.append({"color": "black", "text": run_text, "type": "underline", "_open": False})
                    elif run.font.strike or run.font.double_strike:
                        segments.append({"color": "black", "text": run_text, "type": "strikeout", "_open": False})

        # Close the last open group
        if segments and segments[-1].get("_open"):
            segments[-1]["_open"] = False

        for seg in segments:
            text = seg["text"].strip()
            if not text:
                continue
            highlights.append(
                Highlight(
                    id=str(uuid.uuid4()),
                    source_format="docx",
                    annotation_type=seg["type"],
                    location=location,
                    color=seg["color"],
                    text=text,
                    note=None,
                    timestamp=None,
                    order=order,
                )
            )
            order += 1

    # 3. Append comments as annotations
    for c_id, meta in comments_meta.items():
        text = "".join(meta["text_parts"]).strip()
        if not text and not meta["note"]:
            continue
        # If no text was highlighted, fallback to something
        if not text:
            text = "Attached Note"
            
        highlights.append(
            Highlight(
                id=str(uuid.uuid4()),
                source_format="docx",
                annotation_type="note",
                location=meta["location"] or "Unknown",
                color="yellow",  # Default for comments
                text=text,
                note=meta["note"],
                timestamp=meta["date"],
                order=order,
            )
        )
        order += 1

    return highlights
