"""
exporters.py — Non-AI export formats.

Each exporter is a standalone function:
  export_markdown(highlights, template="simple") -> str
  export_anki(highlights, template="simple")     -> str
  export_notion_md(highlights, template="simple") -> str
  export_pdf(highlights, template="simple") -> bytes
  export_docx(highlights, template="simple") -> bytes
"""

from __future__ import annotations

from .schema import Highlight
from collections import OrderedDict


def _color_label(color: str) -> str:
    return color.replace("_", " ").title()

def _format_meta(h: Highlight) -> str:
    """Helper to format the type and color label"""
    parts = []
    if getattr(h, "annotation_type", "highlight") != "highlight":
        parts.append(getattr(h, "annotation_type").title())
    parts.append(_color_label(h.color))
    return " | ".join(parts)


def export_markdown(highlights: list[Highlight], template: str = "simple") -> str:
    if not highlights:
        return "_No highlights found._\n"

    lines: list[str] = ["# Highlights\n"]
    
    if template == "reading_notes":
        groups: OrderedDict[str, list[Highlight]] = OrderedDict()
        for h in sorted(highlights, key=lambda h: h.order):
            groups.setdefault(h.location, []).append(h)

        for location, items in groups.items():
            lines.append(f"## {location.title()}\n")
            for h in items:
                meta = _format_meta(h)
                lines.append(f"> **[{meta}]** {h.text}")
                if h.note:
                    lines.append(f"> 💬 _{h.note}_")
                if getattr(h, "timestamp", None):
                    lines.append(f"> 🕒 {h.timestamp}")
                lines.append("")
            lines.append("---\n")
    else:
        # Simple flat list
        for h in sorted(highlights, key=lambda h: h.order):
            meta = _format_meta(h)
            lines.append(f"- **[{meta}]** ({h.location}) {h.text}")
            if h.note:
                lines.append(f"  > 💬 {h.note}")
            if getattr(h, "timestamp", None):
                lines.append(f"  > 🕒 {h.timestamp}")
            lines.append("")

    return "\n".join(lines)


_TERM_COLORS = {"yellow", "cyan", "turquoise"}

def export_anki(highlights: list[Highlight], template: str = "simple") -> str:
    if not highlights:
        return ""
    rows: list[str] = []
    for h in sorted(highlights, key=lambda h: h.order):
        meta = _format_meta(h)
        front = h.text.replace("\t", " ").replace("\n", " ")
        if h.color.lower() in _TERM_COLORS:
            back = f"({h.location})"
        else:
            back = f"[{meta}] — {h.location}"
        if h.note:
            back += f" | Note: {h.note}"
        rows.append(f"{front}\t{back}")
    return "\n".join(rows)


def export_notion_md(highlights: list[Highlight], template: str = "simple") -> str:
    if not highlights:
        return "_No highlights found._\n"

    COLOR_EMOJI = {
        "yellow": "🟡", "green": "🟢", "cyan": "🔵", "turquoise": "🩵",
        "pink": "🩷", "red": "🔴", "dark_red": "🔴", "blue": "🔵",
        "dark_blue": "🔵", "violet": "🟣", "teal": "🩵",
    }

    lines: list[str] = ["# Highlights\n"]
    if template == "reading_notes":
        groups: OrderedDict[str, list[Highlight]] = OrderedDict()
        for h in sorted(highlights, key=lambda h: h.order):
            groups.setdefault(h.location, []).append(h)
        for location, items in groups.items():
            lines.append(f"## {location.title()}\n")
            for h in items:
                label = _color_label(h.color)
                emoji = COLOR_EMOJI.get(label.lower(), "📌")
                meta = _format_meta(h)
                lines.append(f"> {emoji} **[{meta}]** {h.text}")
                if h.note:
                    lines.append(f"> 💬 _{h.note}_")
            lines.append("")
    else:
        for h in sorted(highlights, key=lambda h: h.order):
            label = _color_label(h.color)
            emoji = COLOR_EMOJI.get(label.lower(), "📌")
            meta = _format_meta(h)
            lines.append(f"> {emoji} **[{meta}]** ({h.location}) {h.text}")
            if h.note:
                lines.append(f"> 💬 _{h.note}_")
            lines.append("")

    return "\n".join(lines)


def export_pdf(highlights: list[Highlight], template: str = "simple") -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    
    if not highlights:
        pdf.set_font("helvetica", size=12)
        pdf.cell(text="No highlights found.")
        return bytes(pdf.output())

    pdf.set_font("helvetica", style="B", size=16)
    pdf.cell(text="Highlights", new_x="LMARGIN", new_y="NEXT", align="L")
    pdf.ln(5)

    if template == "reading_notes":
        groups: OrderedDict[str, list[Highlight]] = OrderedDict()
        for h in sorted(highlights, key=lambda h: h.order):
            groups.setdefault(h.location, []).append(h)
        for location, items in groups.items():
            pdf.set_font("helvetica", style="B", size=14)
            pdf.cell(text=location.title(), new_x="LMARGIN", new_y="NEXT", align="L")
            pdf.ln(2)
            for h in items:
                meta = _format_meta(h)
                pdf.set_font("helvetica", style="B", size=11)
                pdf.cell(text=f"[{meta}] ", align="L")
                pdf.set_font("helvetica", size=11)
                safe_text = h.text.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(w=0, text=safe_text, align="L", new_x="LMARGIN", new_y="NEXT")
                if h.note:
                    safe_note = h.note.encode('latin-1', 'replace').decode('latin-1')
                    pdf.set_font("helvetica", style="I", size=10)
                    pdf.multi_cell(w=0, text=f"Note: {safe_note}", align="L", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)
            pdf.ln(5)
    else:
        for h in sorted(highlights, key=lambda h: h.order):
            meta = _format_meta(h)
            pdf.set_font("helvetica", style="B", size=11)
            pdf.cell(text=f"[{meta}] ({h.location}) ", align="L")
            pdf.set_font("helvetica", size=11)
            safe_text = h.text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(w=0, text=safe_text, align="L", new_x="LMARGIN", new_y="NEXT")
            if h.note:
                safe_note = h.note.encode('latin-1', 'replace').decode('latin-1')
                pdf.set_font("helvetica", style="I", size=10)
                pdf.multi_cell(w=0, text=f"Note: {safe_note}", align="L", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

    return bytes(pdf.output())


def export_docx(highlights: list[Highlight], template: str = "simple") -> bytes:
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("Highlights", 0)

    if not highlights:
        doc.add_paragraph("No highlights found.")
    else:
        if template == "reading_notes":
            groups: OrderedDict[str, list[Highlight]] = OrderedDict()
            for h in sorted(highlights, key=lambda h: h.order):
                groups.setdefault(h.location, []).append(h)
            for location, items in groups.items():
                doc.add_heading(location.title(), level=1)
                for h in items:
                    meta = _format_meta(h)
                    p = doc.add_paragraph()
                    p.add_run(f"[{meta}] ").bold = True
                    p.add_run(h.text)
                    if h.note:
                        p_note = doc.add_paragraph()
                        p_note.add_run(f"Note: {h.note}").italic = True
        else:
            for h in sorted(highlights, key=lambda h: h.order):
                meta = _format_meta(h)
                p = doc.add_paragraph()
                p.add_run(f"[{meta}] ({h.location}) ").bold = True
                p.add_run(h.text)
                if h.note:
                    p_note = doc.add_paragraph()
                    p_note.add_run(f"Note: {h.note}").italic = True

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
